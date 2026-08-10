from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Read the resume text
with open('M_Azam_Tailored_Resume_DeputyManager_Planning.txt', 'r', encoding='utf-8') as f:
    resume_text = f.read()

with open('M_Azam_Cover_Letter_DeputyManager_Planning.txt', 'r', encoding='utf-8') as f:
    cover_text = f.read()

# Create Word document
doc = Document()

# Style settings
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x25, 0x2C)

# ========== RESUME ==========
# Header
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header.add_run('MOHAMMAD AZAM ALEEM')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact.add_run('Karachi, Pakistan | +92 327 2310 769 | +92 333 3045 707\nazam_aleem@hotmail.com | linkedin.com/in/muhammed-azam-aleem-7a35ba18')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Resume title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RESUME — DEPUTY MANAGER PLANNING')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)

doc.add_paragraph()  # spacing

# Process resume text
for line in resume_text.split('\n'):
    line = line.strip()
    if not line:
        doc.add_paragraph()
        continue
    
    # Skip decorative lines
    if line.startswith('══') or line.startswith('━━') or line.startswith('──'):
        continue
    
    # Section headers (all caps)
    if line.isupper() and len(line) > 5 and not line.startswith('▸') and not line.startswith('✅'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        continue
    
    # Bullet points
    if line.startswith('▸') or line.startswith('✅') or line.startswith('•'):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(line)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        continue
    
    # Regular text
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

# Page break before cover letter
doc.add_page_break()

# ========== COVER LETTER ==========
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run('COVER LETTER — DEPUTY MANAGER PLANNING')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)

doc.add_paragraph()  # spacing

# Process cover letter text
for line in cover_text.split('\n'):
    line = line.strip()
    if not line:
        doc.add_paragraph()
        continue
    
    # Section headers with emoji
    if line.startswith('🔹'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        continue
    
    # Bullet points
    if line.startswith('•') or line.startswith('✅'):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(line)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        continue
    
    # Regular text
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

# Save
output_path = os.path.join(os.getcwd(), 'M_Azam_Complete_Application_DeputyManager_Planning.docx')
doc.save(output_path)
print(f"Word document created: {output_path}")